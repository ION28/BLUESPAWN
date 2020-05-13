from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.urls import reverse
from django.views import generic
from django import template
from django.core.files.storage import FileSystemStorage
from django.conf import settings

from lxml import etree as ET
import uuid
import os
from docx import Document

def generate_report(scanxml):
    report_root = ET.parse(scanxml).getroot()

    report = Document()
    report.add_heading('BLUESPAWN Host Compromise Analysis Report (HCAR)', 0)

    # Parse XML File
    detections = list(report_root.findall('hunt/detection'))
    detections_hunts = []
    detections_files = []
    detections_registry = []
    detections_processes = []
    detections_eventlogs = []
    file_cnt = 0
    registry_cnt = 0
    process_cnt = 0
    eventlog_cnt = 0

    for hunt in report_root.findall('hunt'):
        if hunt.find('detection') is not None:
            detections_hunts.append([hunt.find('name').text])

    for detection in detections:
        if detection.get('type') == 'File':
            file_cnt += 1
            detections_files.append([
                detection.find('path').text if detection.find('path') is not None else "N/A", 
                detection.find('size').text if detection.find('size') is not None else "N/A", 
                detection.find('md5').text if detection.find('md5') is not None else "N/A", 
                detection.find('sha1').text if detection.find('sha1') is not None else "N/A", 
                detection.find('sha256').text if detection.find('sha256') is not None else "N/A",
                detection.getparent().find('name').text if detection.getparent().find('name').text is not None else "N/A"
            ])
        elif detection.get('type') == 'Registry':
            registry_cnt += 1
            detections_registry.append([
                detection.find('key').text if detection.find('key') is not None else "N/A", 
                detection.find('value').text if detection.find('value') is not None else "N/A", 
                detection.find('data').text if detection.find('data') is not None else "N/A",
                detection.getparent().find('name').text if detection.getparent().find('name').text is not None else "N/A"
            ]) 
        elif detection.get('type') == 'Process':
            process_cnt += 1
            detections_processes.append([
                detection.find('path').text if detection.find('path') is not None else "N/A", 
                detection.find('cmdline').text if detection.find('cmdline') is not None else "N/A", 
                detection.find('username').text if detection.find('username') is not None else "N/A",
                detection.find('pid').text if detection.find('pid') is not None else "N/A",
                detection.find('method').text if detection.find('method') is not None else "N/A"
            ]) 
        '''
        elif detection.get('type') == 'Event':
            eventlog_cnt += 1
            detections_eventlogs.append([
                detection.find('id').text if detection.find('id') is not None else "N/A", 
                detection.find('channel').text if detection.find('channel') is not None else "N/A"
            ]) 
        '''

    # MITRE ATT&CK IDs Table
    if file_cnt > 0 or registry_cnt > 0 or process_cnt > 0 or eventlog_cnt > 0:
        report.add_heading('MTIRE ATT&CK IDs:', level=1)
        mitre_table = report.add_table(rows=1, cols=3, style='Table Grid')
        mitre_table_header = mitre_table.rows[0].cells
        mitre_table_header[0].text = 'Technique ID/Name:'
        mitre_table_header[1].text = 'Tactic:'
        mitre_table_header[2].text = 'Description:'
        for item in detections_hunts:
            r = mitre_table.add_row().cells
            r[0].text = str(item[0])

    # Indicators of Attack
    report.add_heading('Indicators of Attack (IOAs)', level=1)
    report.add_heading('Files (%i)' % file_cnt, level=2)
    if file_cnt > 0:
        file_table = report.add_table(rows=1, cols=3, style='Table Grid')
        file_table_header = file_table.rows[0].cells
        file_table_header[0].text = 'Filename:'
        file_table_header[1].text = 'Details:'
        file_table_header[2].text = 'Associated Technique:'
        for item in detections_files:
            r = file_table.add_row().cells
            r[0].text = str(item[0])
            r[1].text = ('Size: ' + str(item[1]) + '\nMD5: ' + str(item[2]) + 
                '\nSHA1: ' + str(item[3]) + '\nSHA256: ' + str(item[4]))
            r[2].text = str(item[5])

    report.add_heading('Registry (%i)' % registry_cnt, level=2)
    if registry_cnt > 0:
        registry_table = report.add_table(rows=1, cols=4, style='Table Grid')
        registry_table_header = registry_table.rows[0].cells
        registry_table_header[0].text = 'Key:'
        registry_table_header[1].text = 'Value:'
        registry_table_header[2].text = 'Data:'
        registry_table_header[3].text = 'Associated Technique:'
        for item in detections_registry:
            r = registry_table.add_row().cells
            r[0].text = str(item[0])
            r[1].text = str(item[1])
            r[2].text = str(item[2])
            r[3].text = str(item[3])
    
    report.add_heading('Processes (%i)' % process_cnt, level=2)
    if process_cnt > 0:
        process_table = report.add_table(rows=1, cols=5, style='Table Grid')
        process_table_header = process_table.rows[0].cells
        process_table_header[0].text = 'Path:'
        process_table_header[1].text = 'Command Line:'
        process_table_header[2].text = 'Username:'
        process_table_header[3].text = 'PID:'
        process_table_header[4].text = 'Method:'
        for item in detections_processes:
            r = process_table.add_row().cells
            r[0].text = str(item[0])
            r[1].text = str(item[1])
            r[2].text = str(item[2])
            r[3].text = str(item[3])
            r[4].text = str(item[4])
    '''    
    report.add_heading('Event Logs (%i)' % eventlog_cnt, level=2)
    if eventlog_cnt > 0:
        eventlog_table = report.add_table(rows=1, cols=3, style='Table Grid')
        eventlog_table_header = eventlog_table.rows[0].cells
        eventlog_table_header[0].text = 'ID:'
        eventlog_table_header[1].text = 'Channel:'
        eventlog_table_header[2].text = 'Information:'
        for item in detections_eventlogs:
            r = eventlog_table.add_row().cells
            r[0].text = str(item[0])
            r[1].text = str(item[1])
    ''' 

    report_name = str(uuid.uuid4()) + '.docx'
    report.save(os.path.join(getattr(settings, 'BASE_DIR'), 'media/scans/' + report_name))

    return report_name

def UploadScanView(request):
    if request.method == 'POST' and request.FILES['scanxml']:
        scan = request.FILES['scanxml']
        fs = FileSystemStorage()
        ext = scan.name.split('.')[-1]
        filename = "%s.%s" % (uuid.uuid4(), ext)

        if ext != 'xml':
            return redirect('/report?error=1')

        upload = fs.save('scans/' + filename, scan)

        try:
            report_name = generate_report(os.path.join(getattr(settings, 'BASE_DIR'), 'media/scans/' + filename))
            fs.delete(os.path.join(getattr(settings, 'BASE_DIR'), 'media/scans/' + filename))
            return render(request, 'report.html', {
                'generated_report_url' : report_name
            })
        except:
            return redirect('/report?error=1')

    return render(request, 'report.html')

